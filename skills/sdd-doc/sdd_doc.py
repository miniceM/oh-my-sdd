#!/usr/bin/env python3
"""
sdd_doc.py —— /sdd-doc 命令核心脚本

读取 openspec change 的 spec 产物（proposal.md + specs/*.md + .meta.json），
用 pandoc 生成正文 docx，再用 python-docx 以对象模型填充企业模板，
输出符合企业格式的 Word 文档。

技术路线（demo1/3/3b 验证）：
  pandoc 生成正文 markdown → docx + python-docx 对象级填充模板
模板始终保持完整，不做 XML 切分拼接（demo2 已证伪该路线）。

用法：
  python3 sdd_doc.py [change-slug] [--project NAME] [--clean] [--stage spec]
"""

from __future__ import annotations

import argparse
import json
import os
import platform
import re
import shutil
import subprocess
import sys
import tempfile
from copy import deepcopy
from datetime import datetime
from pathlib import Path

# 延迟 import：python-docx 缺失时由 check_python_deps() 给友好提示，
# 而非顶部 import 直接抛 ModuleNotFoundError。
try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    _DOCX_AVAILABLE = True
except ImportError:
    Document = None  # type: ignore[assignment]
    OxmlElement = None  # type: ignore[assignment]
    qn = None  # type: ignore[assignment]
    _DOCX_AVAILABLE = False


# ============================================================
# 常量
# ============================================================

SCRIPT_DIR = Path(__file__).resolve().parent
BUILTIN_TEMPLATE = SCRIPT_DIR / "templates" / "reference.docx"
OUTPUT_FILENAME_SUFFIX = "需求规格说明书.docx"

# 封面/信息表/修订表用到的模板样式 styleId
STYLE_BODY_INDENT = "28"   # 正文文本首行缩进
STYLE_HEADING1 = "2"
STYLE_HEADING2 = "3"
STYLE_HEADING3 = "4"
STYLE_HEADING4 = "5"


# ============================================================
# 错误处理
# ============================================================

class SddDocError(Exception):
    """用户可见错误，消息直接打印。"""


# ============================================================
# 前置检查
# ============================================================

def check_pandoc() -> None:
    """检查 pandoc 可用，缺失则按 OS 给安装命令并退出。"""
    if shutil.which("pandoc"):
        return
    system = platform.system()
    hints = {
        "Darwin": "brew install pandoc",
        "Windows": "winget install --id JohnMacFarlane.Pandoc",
        "Linux": "apt install pandoc  # 或 yum install pandoc",
    }
    cmd = hints.get(system, "从 https://pandoc.org/installing.html 安装")
    raise SddDocError(
        f"未找到 pandoc。/sdd-doc 依赖 pandoc 做正文转换。\n"
        f"请先安装（{system}）：  {cmd}"
    )


def check_python_deps() -> None:
    """检查 python-docx + lxml。"""
    if not _DOCX_AVAILABLE:
        raise SddDocError(
            "缺少 Python 依赖 python-docx（或 lxml）。\n"
            "请运行：pip install python-docx lxml"
        )


# ============================================================
# 模板查找（三层优先级）
# ============================================================

def find_template(change_dir: Path, project_root: Path) -> Path:
    """三层优先级查找模板：
    1. change 本地：openspec/changes/<slug>/reference.docx
    2. 项目配置：<project-root>/.sdd-doc-template（单行路径文件）
    3. plugin 内置：skills/sdd-doc/templates/reference.docx
    """
    # 1. change 本地覆盖
    local = change_dir / "reference.docx"
    if local.is_file():
        return local

    # 2. 项目级配置（读 .sdd-doc-template 单行路径文件）
    cfg_marker = project_root / ".sdd-doc-template"
    if cfg_marker.is_file():
        tpl_path = cfg_marker.read_text(encoding="utf-8").strip()
        tpl = Path(tpl_path)
        if not tpl.is_absolute():
            tpl = project_root / tpl
        if tpl.is_file():
            return tpl

    # 3. 内置默认
    if BUILTIN_TEMPLATE.is_file():
        return BUILTIN_TEMPLATE

    raise SddDocError(
        f"未找到 Word 模板。查找顺序：\n"
        f"  1. {local}\n"
        f"  2. {cfg_marker} 指向的路径\n"
        f"  3. 内置 {BUILTIN_TEMPLATE}\n"
        f"请放置模板文件或配置 .sdd-doc-template。"
    )


def load_capability_names(project_root: Path, cap_names_arg: str | None) -> dict:
    """加载 capability 显示名映射。

    来源优先级：命令行 --capability-names > 项目 .sdd-doc-names 配置文件。
    格式：auth=认证,user=用户管理（逗号分隔）。
    """
    pairs: list[str] = []
    # 1. 项目配置文件
    cfg = project_root / ".sdd-doc-names"
    if cfg.is_file():
        pairs.extend(cfg.read_text(encoding="utf-8").split(","))

    # 2. 命令行参数（覆盖配置）
    if cap_names_arg:
        pairs.extend(cap_names_arg.split(","))

    names: dict[str, str] = {}
    for pair in pairs:
        pair = pair.strip()
        if not pair or "=" not in pair:
            continue
        k, v = pair.split("=", 1)
        names[k.strip()] = v.strip()
    return names


# ============================================================
# spec 产物读取
# ============================================================

def resolve_change_dir(slug: str, project_root: Path) -> Path:
    change_dir = project_root / "openspec" / "changes" / slug
    if not change_dir.is_dir():
        raise SddDocError(
            f"未找到 change 目录：{change_dir}\n"
            f"请确认 slug 正确，或先运行 /sdd-spec 创建规格。"
        )
    return change_dir


def read_meta(change_dir: Path) -> dict:
    meta_path = change_dir / ".meta.json"
    if not meta_path.is_file():
        return {}
    try:
        return json.loads(meta_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as e:
        raise SddDocError(f".meta.json 解析失败：{e}")


def collect_markdown(change_dir: Path) -> str:
    """合并 proposal.md + specs/*.md 为单一 markdown 流。"""
    parts = []
    proposal = change_dir / "proposal.md"
    if proposal.is_file():
        parts.append(proposal.read_text(encoding="utf-8"))

    specs_dir = change_dir / "specs"
    if specs_dir.is_dir():
        for spec_file in sorted(specs_dir.rglob("*.md")):
            parts.append(f"\n\n---\n\n{spec_file.read_text(encoding='utf-8')}")

    if not parts:
        raise SddDocError(
            f"在 {change_dir} 下未找到 proposal.md 或 specs/*.md，无法生成正文。"
        )
    return "\n".join(parts)


# ============================================================
# spec/proposal 结构化解析（v2 章节映射基础）
# ============================================================

def _normalize_proposal_heading(heading: str) -> str:
    """把 proposal 的 ## 标题标准化为固定 key。"""
    h = heading.strip()
    if "背景" in h or "目标" in h or "目的" in h:
        return "background"
    if "范围" in h:
        return "scope"
    if "验收" in h:
        return "acceptance"
    return h  # 非标准标题保留原文作 key


def parse_proposal(md: str) -> dict:
    """解析 proposal.md 为结构化数据，按 ## 分块，标题标准化。

    返回 {title, background, scope, acceptance, ...}。
    无标准 ## 标题时整段归 "intro"（容错）。
    """
    lines = md.splitlines()
    title = ""
    for line in lines:
        if line.startswith("# "):
            title = line[2:].strip()
            break

    blocks: dict[str, str] = {}
    current_key = "intro"
    current_lines: list[str] = []
    for line in lines:
        m = re.match(r"^##\s+(.+?)\s*$", line)
        if m:
            blocks[current_key] = "\n".join(current_lines).strip()
            current_key = _normalize_proposal_heading(m.group(1))
            current_lines = []
        elif line.startswith("# "):
            continue  # 跳过文档级标题
        else:
            current_lines.append(line)
    blocks[current_key] = "\n".join(current_lines).strip()

    # 清理空的 intro（通常只是文档标题前的空白）
    if not blocks.get("intro"):
        blocks.pop("intro", None)

    blocks["title"] = title
    return blocks


def _parse_delta(md: str) -> list[dict]:
    """解析 openspec delta 格式的 spec.md，返回 requirement 列表。

    每项: {name, desc, scenarios: [{name, steps: [{type, text}]}]}
    依赖固定格式：### Requirement: / #### Scenario: / - **WHEN|THEN|AND|BUT** ...
    """
    requirements: list[dict] = []
    current_req: dict | None = None
    current_scenario: dict | None = None
    in_req_desc = False

    for line in md.splitlines():
        m_req = re.match(r"^###\s+Requirement:\s*(.+?)\s*$", line)
        m_scn = re.match(r"^####\s+Scenario:\s*(.+?)\s*$", line)
        m_step = re.match(r"^-\s+\*\*(WHEN|THEN|AND|BUT)\*\*\s*(.*)$", line)

        if m_req:
            if current_req:
                requirements.append(current_req)
            current_req = {"name": m_req.group(1), "desc": "", "scenarios": []}
            current_scenario = None
            in_req_desc = True
        elif m_scn:
            if current_req is None:
                current_req = {"name": "", "desc": "", "scenarios": []}
            current_scenario = {"name": m_scn.group(1), "steps": []}
            current_req["scenarios"].append(current_scenario)
            in_req_desc = False
        elif m_step:
            if current_scenario is not None:
                current_scenario["steps"].append({
                    "type": m_step.group(1),
                    "text": m_step.group(2).strip(),
                })
            in_req_desc = False
        else:
            # requirement 标题下、scenario 前的描述行
            if in_req_desc and current_req is not None and line.strip():
                current_req["desc"] = (
                    current_req["desc"] + "\n" + line.strip()
                    if current_req["desc"] else line.strip()
                )

    if current_req:
        requirements.append(current_req)
    return requirements


def parse_specs(change_dir: Path) -> list[dict]:
    """解析 specs/<capability>/spec.md，返回 capability 列表。

    每项: {name, requirements: [{name, desc, scenarios: [...]}]}
    """
    specs_dir = change_dir / "specs"
    if not specs_dir.is_dir():
        return []

    capabilities: list[dict] = []
    for cap_dir in sorted(specs_dir.iterdir()):
        if not cap_dir.is_dir():
            continue
        cap: dict = {"name": cap_dir.name, "requirements": []}
        for spec_file in sorted(cap_dir.glob("*.md")):
            md = spec_file.read_text(encoding="utf-8")
            cap["requirements"].extend(_parse_delta(md))
        capabilities.append(cap)
    return capabilities


# ============================================================
# 元数据抽取
# ============================================================

def git_user_name(project_root: Path) -> str:
    try:
        out = subprocess.run(
            ["git", "config", "user.name"],
            cwd=str(project_root), capture_output=True, text=True, timeout=5,
        )
        if out.returncode == 0 and out.stdout.strip():
            return out.stdout.strip()
    except Exception:
        pass
    return ""


def extract_metadata(
    change_dir: Path, project_root: Path, meta: dict, project_arg: str | None
) -> dict:
    """按优先级抽取封面/信息表元数据。"""
    created_at = meta.get("created_at", "")
    try:
        date_str = datetime.fromisoformat(created_at.replace("Z", "+00:00")).strftime("%Y-%m-%d")
    except Exception:
        date_str = datetime.now().strftime("%Y-%m-%d")

    project_title = (
        project_arg
        or meta.get("project_title")
        or f"{{{{project_title}}}}"  # 保留模板占位符，让用户手填
    )

    return {
        "project_title": project_title,
        "version": meta.get("version", "V1.0"),
        "author": meta.get("author") or git_user_name(project_root) or "",
        "date": date_str,
        "department": meta.get("department", "科技部"),
        "control_level": "内部资料",
    }


def extract_revisions(change_dir: Path, project_root: Path, meta: dict) -> list[dict]:
    """抽取修订记录行。每行：{version, date, summary, section, author, reviewer}

    方案 C：单行初始记录。每次生成填一行"初始创建"，后续修订由用户在 Word 里手填。
    简单稳定，适合 MVP；后续若需自动追踪 git 历史，可切换到方案 A。
    """
    version = meta.get("version", "V1.0")
    try:
        created = meta.get("created_at", "")
        date = datetime.fromisoformat(created.replace("Z", "+00:00")).strftime("%Y-%m-%d")
    except Exception:
        date = datetime.now().strftime("%Y-%m-%d")
    return [{
        "version": version,
        "date": date,
        "summary": "初始创建",
        "section": "全文",
        "author": meta.get("author") or git_user_name(project_root) or "",
        "reviewer": "",
    }]


# ============================================================
# pandoc 正文生成
# ============================================================

def generate_body_docx(
    markdown: str, template_path: Path, work_dir: Path
) -> Path:
    """用 pandoc 把 markdown 转成正文 docx（套用模板样式 + TOC）。

    返回临时 docx 路径。后续用 python-docx 抽取其 body 段落。
    """
    md_file = work_dir / "_body.md"
    md_file.write_text(markdown, encoding="utf-8")
    out_file = work_dir / "_body.docx"

    cmd = [
        "pandoc", str(md_file),
        "--reference-doc", str(template_path),
        "--toc", "--toc-depth=4",
        "-o", str(out_file),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise SddDocError(
            f"pandoc 转换失败（exit {result.returncode}）：\n{result.stderr}"
        )
    return out_file


# ============================================================
# 模板填充（python-docx 对象级操作）
# ============================================================

def _set_para_text(p, new_text: str) -> None:
    """改段落文字，保留首个 run 的格式。"""
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(new_text)


def _set_cell_text(cell, new_text: str) -> None:
    """改单元格文字，保留首段首 run 格式。"""
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = new_text
    else:
        cell.paragraphs[0].add_run(new_text)


def fill_cover(doc: Document, data: dict) -> None:
    """填充封面：项目标题（含 {{project_title}} 占位的段落）+ 版本号。"""
    for p in doc.paragraphs[:20]:
        if "{{project_title}}" in p.text:
            _set_para_text(p, data["project_title"])
        elif p.text.strip() == "Vn.n":
            _set_para_text(p, data["version"])


def fill_info_table(doc: Document, data: dict) -> None:
    """填充文档信息表（编写人/发布日期/控制级别/制定部门）。

    遍历表格，按左列标签匹配右列填值。
    """
    info = {
        "编写人": data["author"],
        "发布日期": data["date"],
        "控制级别": data["control_level"],
        "制定部门": data["department"],
    }
    for table in doc.tables:
        if len(table.rows) < 1 or len(table.columns) < 2:
            continue
        labels = [r.cells[0].paragraphs[0].text.strip() for r in table.rows]
        if "编写人" in labels:
            for row in table.rows:
                label = row.cells[0].paragraphs[0].text.strip()
                if label in info:
                    _set_cell_text(row.cells[1], info[label])
            return


def fill_revisions_table(doc: Document, revisions: list[dict]) -> None:
    """填充修订记录表。

    表头识别：含"版本""时间""变更概要""作者"的行。
    数据从 revisions list 依次填入表头之后的行；行不够时追加新行。
    """
    header_keys = ["版本", "时间", "变更概要", "修改章节", "作者", "审核"]
    rev_table = None
    for table in doc.tables:
        if len(table.rows) < 1:
            continue
        header = [c.paragraphs[0].text.strip() for c in table.rows[0].cells]
        if "版本" in header and "变更概要" in header:
            rev_table = table
            break

    if rev_table is None:
        return  # 模板无修订表，跳过

    # 建立列索引（按表头文字匹配，容错列顺序）
    col_map = {}
    for ci, cell in enumerate(rev_table.rows[0].cells):
        label = cell.paragraphs[0].text.strip()
        if label in header_keys:
            col_map[label] = ci

    # 数据行从第 2 行（index=1）开始填
    for ri, rev in enumerate(revisions):
        row_idx = ri + 1
        # 行不够则追加（复用上一行结构）
        if row_idx >= len(rev_table.rows):
            new_row = deepcopy(rev_table.rows[-1]._tr)
            rev_table._tbl.append(new_row)
        row = rev_table.rows[row_idx]
        for label, ci in col_map.items():
            key = {
                "版本": "version", "时间": "date", "变更概要": "summary",
                "修改章节": "section", "作者": "author", "审核": "reviewer",
            }[label]
            _set_cell_text(row.cells[ci], str(rev.get(key, "")))


def insert_body(doc: Document, body_docx: Path, clean: bool) -> None:
    """v1 扁平插入（已弃用，由 map_to_template 取代）。保留签名供向后兼容，但不再被 generate 调用。"""
    return  # no-op；v2 使用 map_to_template


# ============================================================
# v2 章节映射填充（map_to_template 及其辅助函数）
# ============================================================

# 模板功能块下可自动填充的子章节（按标题文字匹配）
_SUBSECTION_OVERVIEW = "功能概述"
_SUBSECTION_LOGIC = "处理逻辑"


def _get_pstyle(element) -> str | None:
    """从 w:p 元素取 pStyle 的 w:val，非段落或无样式返回 None。"""
    if element.tag != qn("w:p"):
        return None
    ppr = element.find(qn("w:pPr"))
    if ppr is None:
        return None
    pstyle = ppr.find(qn("w:pStyle"))
    if pstyle is None:
        return None
    return pstyle.get(qn("w:val"))


def _get_text(element) -> str:
    """取 w:p 元素的所有 w:t 拼接文本。"""
    if element.tag != qn("w:p"):
        return ""
    return "".join((t.text or "") for t in element.iter(qn("w:t")))


def _make_heading_para(text: str, style_id: str):
    """构造一个标题段落元素（带 pStyle）。"""
    p = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    pstyle = OxmlElement("w:pStyle")
    pstyle.set(qn("w:val"), style_id)
    ppr.append(pstyle)
    p.append(ppr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def _make_body_para(text: str, style_id: str = STYLE_BODY_INDENT):
    """构造正文段落（默认 styleId=28 正文文本首行缩进）。"""
    return _make_heading_para(text, style_id)


def _make_body_paras_from_text(text: str) -> list:
    """多行文本 → 段落元素列表（每非空行一段）。空文本给占位。"""
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if not lines:
        return [_make_body_para("（待补充）")]
    return [_make_body_para(ln) for ln in lines]


def _find_heading1_by_text(doc: Document, keyword: str):
    """找首个含 keyword 的 heading1 段落元素，找不到返回 None。"""
    for child in doc.element.body:
        if _get_pstyle(child) == STYLE_HEADING1 and keyword in _get_text(child):
            return child
    return None


def _find_first_heading1(doc: Document):
    """找首个非空 heading1 段落元素。"""
    for child in doc.element.body:
        if _get_pstyle(child) == STYLE_HEADING1 and _get_text(child).strip():
            return child
    return None


def _set_element_text(element, new_text: str) -> None:
    """改 w:p 元素的文本（首个 w:t 写入，其余清空；无 w:t 则追加 run）。保留段落格式。"""
    ts = list(element.iter(qn("w:t")))
    if ts:
        ts[0].text = new_text
        ts[0].set(qn("xml:space"), "preserve")
        for t in ts[1:]:
            t.text = ""
    else:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = new_text
        r.append(t)
        element.append(r)


def _renumber_template_heading1s(doc: Document) -> None:
    """给模板原有 L1 加编号前缀，与新增'1 总体说明'保持编号一致。

    '功能详细分析及设计' → '2 功能详细分析及设计'
    '非功能需求分析' → '3 非功能需求分析'
    幂等：已带目标前缀则跳过。
    """
    mapping = {"功能详细分析及设计": "2 ", "非功能需求分析": "3 "}
    for child in doc.element.body:
        if _get_pstyle(child) != STYLE_HEADING1:
            continue
        text = _get_text(child).strip()
        if text in mapping:
            _set_element_text(child, mapping[text] + text)


def _find_functional_region(doc: Document) -> tuple[int | None, int]:
    """定位功能区：'功能详细分析及设计'h1 到下一个 h1（通常'非功能需求分析'）之间。

    返回 (func_h1_idx, next_h1_idx)。基于文本匹配避免空 heading 占位干扰。
    """
    children = list(doc.element.body)
    func_h1_idx: int | None = None
    for i, child in enumerate(children):
        if (_get_pstyle(child) == STYLE_HEADING1
                and "功能" in _get_text(child)
                and "设计" in _get_text(child)):
            func_h1_idx = i
            break
    if func_h1_idx is None:
        return None, 0
    next_h1_idx = len(children)
    for i in range(func_h1_idx + 1, len(children)):
        if _get_pstyle(children[i]) == STYLE_HEADING1 and _get_text(children[i]).strip():
            next_h1_idx = i
            break
    return func_h1_idx, next_h1_idx


def find_function_block_range(doc: Document) -> tuple[int | None, int]:
    """功能区内首个【有文本】的 heading2 到下一个 heading2 之间（body 子元素索引）。

    跳过模板里的空 heading2 格式占位，确保圈到真实功能块。
    """
    func_start, func_end = _find_functional_region(doc)
    if func_start is None:
        return None, 0
    children = list(doc.element.body)
    block_start: int | None = None
    for i in range(func_start + 1, func_end):
        if (_get_pstyle(children[i]) == STYLE_HEADING2
                and _get_text(children[i]).strip()):
            block_start = i
            break
    if block_start is None:
        return None, 0
    block_end = func_end
    for i in range(block_start + 1, func_end):
        if (_get_pstyle(children[i]) == STYLE_HEADING2
                and _get_text(children[i]).strip()):
            block_end = i
            break
    return block_start, block_end


def delete_sample_function_blocks(doc: Document) -> int:
    """删除功能区内所有示例功能块（首个 heading2 到下一个 h1 之间）。

    保留功能区总览（"功能详细分析及设计"h1 与首个 heading2 之间的说明段落和总览表）。
    返回删除元素数。
    """
    func_start, func_end = _find_functional_region(doc)
    if func_start is None:
        return 0
    children = list(doc.element.body)
    first_h2: int | None = None
    for i in range(func_start + 1, func_end):
        if _get_pstyle(children[i]) == STYLE_HEADING2 and _get_text(children[i]).strip():
            first_h2 = i
            break
    if first_h2 is None:
        return 0
    body = doc.element.body
    to_remove = children[first_h2:func_end]
    for el in to_remove:
        body.remove(el)
    return len(to_remove)


# ---- 表格单元格填充（v3：保留模板表格，选择性填 spec 数据）----

def _set_cell_el_text(tc_el, text: str) -> None:
    """改单元格文字（支持多行，每行一段，保留首段 pPr 样式）。"""
    lines = text.split("\n") if text else [""]
    paras = tc_el.findall(qn("w:p"))
    if not paras:
        p = OxmlElement("w:p")
        tc_el.append(p)
        paras = [p]
    _set_element_text(paras[0], lines[0])
    for p in paras[1:]:
        tc_el.remove(p)
    first_ppr = paras[0].find(qn("w:pPr"))
    for line in lines[1:]:
        new_p = OxmlElement("w:p")
        if first_ppr is not None:
            new_p.append(deepcopy(first_ppr))
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = line
        r.append(t)
        new_p.append(r)
        tc_el.append(new_p)


def _fill_overview_table(tbl_el, cap: dict) -> None:
    """填功能概述表（5×2）。

    行1 右列"功能名称" ← capability 显示名
    行4 右列"功能说明" ← 首个 requirement 描述（无则占位）
    其余行（功能编号/变更类型/触发入口）保留模板示例作指引。
    """
    rows = tbl_el.findall(qn("w:tr"))
    reqs = cap.get("requirements", [])
    desc = (
        reqs[0]["desc"] if reqs and reqs[0].get("desc")
        else f"（{cap['display_name']} 功能说明待补充）"
    )
    fills = {1: cap["display_name"], 4: desc}
    for ri, text in fills.items():
        if ri < len(rows):
            cells = rows[ri].findall(qn("w:tc"))
            if len(cells) > 1:
                _set_cell_el_text(cells[1], text)


def _format_logic_text(cap: dict) -> str:
    """处理逻辑多行文本：每个 requirement 名 + 描述 + 各 scenario 的 WHEN/THEN。"""
    lines: list[str] = []
    for req in cap.get("requirements", []):
        lines.append(f"■ {req['name']}")
        if req.get("desc"):
            lines.append(req["desc"])
        for scn in req.get("scenarios", []):
            lines.append(f"场景：{scn['name']}")
            steps = scn.get("steps", [])
            if steps:
                lines.append(" ".join(f"{s['type']} {s['text']}" for s in steps))
    return "\n".join(lines) if lines else "（待补充）"


def _fill_logic_table(tbl_el, cap: dict) -> None:
    """填处理逻辑表（1×1）：单元格 = scenario WHEN-THEN 多行文本。"""
    rows = tbl_el.findall(qn("w:tr"))
    if not rows:
        return
    cells = rows[0].findall(qn("w:tc"))
    if cells:
        _set_cell_el_text(cells[0], _format_logic_text(cap))


def fill_function_block_tables(block_elements: list, cap: dict) -> None:
    """在功能块元素列表里，按 h3 标题识别并填充功能概述表、处理逻辑表。

    每个表只填第一个（h3 下首个表格）。其余表格（接口/会计/影响范围等）原样保留。
    """
    current_h3: str | None = None
    filled_overview = False
    filled_logic = False
    for el in block_elements[1:]:  # 跳过首个 heading2
        ps = _get_pstyle(el)
        if ps == STYLE_HEADING3:
            current_h3 = _get_text(el).strip()
        elif el.tag == qn("w:tbl"):
            if current_h3 == _SUBSECTION_OVERVIEW and not filled_overview:
                _fill_overview_table(el, cap)
                filled_overview = True
            elif current_h3 == _SUBSECTION_LOGIC and not filled_logic:
                _fill_logic_table(el, cap)
                filled_logic = True


def insert_overview_section(doc: Document, proposal: dict) -> None:
    """在"功能详细分析及设计"heading1 前插入"1 总体说明"及子章节。

    用 addprevious 逆序插入保证顺序。
    """
    _h1 = _find_heading1_by_text(doc, "功能详细分析及设计")
    func_h1 = _h1 if _h1 is not None else _find_first_heading1(doc)
    if func_h1 is None:
        return  # 模板无 heading1，无法定位，放弃

    sections = [
        ("1 总体说明", STYLE_HEADING1, None),
        ("1.1 项目信息", STYLE_HEADING2, proposal.get("background") or proposal.get("intro") or ""),
        ("1.2 功能列表", STYLE_HEADING2, proposal.get("scope") or ""),
        ("1.3 验收标准", STYLE_HEADING2, proposal.get("acceptance") or ""),
    ]

    elements: list = []
    for title, sid, body in sections:
        elements.append(_make_heading_para(title, sid))
        if body is not None:
            elements.extend(_make_body_paras_from_text(body))

    # 正序 addprevious：每次插到 func_h1 正前，自然形成 A B C <func_h1> 顺序
    for el in elements:
        func_h1.addprevious(el)


def fill_nonfunctional_placeholders(doc: Document) -> int:
    """在"非功能需求分析"heading1 区下，对无内容的 heading3 子章节插入"（待补充）"占位。

    安全增量：只在 heading3 后紧跟下一个标题（说明无内容）时插入，避免重复。
    返回插入数。
    """
    body = doc.element.body
    in_nonfunc = False
    inserted = 0
    for child in list(body):
        ps = _get_pstyle(child)
        text = _get_text(child)
        if ps == STYLE_HEADING1 and text.strip():
            in_nonfunc = "非功能" in text
            continue
        if in_nonfunc and ps == STYLE_HEADING3:
            nxt = child.getnext()
            nxt_ps = _get_pstyle(nxt) if nxt is not None else None
            # 下一元素是标题或无 → 该子章节无内容，插占位
            if nxt is None or nxt_ps in (STYLE_HEADING1, STYLE_HEADING2, STYLE_HEADING3, STYLE_HEADING4):
                child.addnext(_make_body_para("（待补充）"))
                inserted += 1
    return inserted


def map_to_template(
    doc: Document, proposal: dict, capabilities: list[dict], cap_names: dict
) -> None:
    """v3 核心：按章节映射填充，保留模板表格结构。

    1. 提取首个示例功能块（含所有表格）作为模板
    2. 删除示例功能块（保留功能区总览）
    3. 在"功能详细分析及设计"前插入"1 总体说明"（proposal 内容）
    4. 给模板 L1 加编号
    5. 每个 capability：deepcopy 功能块模板 → 改 heading2 标题 → 填功能概述表+处理逻辑表
       → 插入功能区总览之后（其余表格原样保留作指引）
    6. 非功能区无内容子章节插占位
    """
    for cap in capabilities:
        cap["display_name"] = cap_names.get(cap["name"], cap["name"])

    # A. 提取首个示例功能块（含表格）作为 deepcopy 模板
    start, end = find_function_block_range(doc)
    if start is None:
        raise SddDocError(
            "模板未找到功能块（heading2）作为模板源，无法映射 capability。"
        )
    block_template = list(doc.element.body)[start:end]

    # B. 删除示例功能块（保留功能区总览：说明段落 + 总览表）
    deleted = delete_sample_function_blocks(doc)

    # C. 插入"1 总体说明" + 给模板 L1 加编号
    insert_overview_section(doc, proposal)
    _renumber_template_heading1s(doc)

    # D. 定位功能块插入锚点：功能区总览最后一个元素之后
    #    （"功能详细分析及设计"h1 之后、原首个 heading2 之前——示例已删，故为总览末尾）
    _h1 = _find_heading1_by_text(doc, "功能详细分析及设计")
    func_h1 = _h1 if _h1 is not None else _find_first_heading1(doc)
    if func_h1 is None:
        raise SddDocError("模板无'功能详细分析及设计'heading1，无法插入功能块。")
    anchor = func_h1
    nxt = func_h1.getnext()
    # 锚点停在总览末尾：遇任何 heading（h1 如"非功能" 或 h2 首个功能块）即停
    while nxt is not None and _get_pstyle(nxt) not in (STYLE_HEADING1, STYLE_HEADING2):
        anchor = nxt
        nxt = nxt.getnext()

    # E. 每个 capability：deepcopy 功能块模板，改标题，填两个表，顺序插入
    for cap in capabilities:
        new_block = [deepcopy(el) for el in block_template]
        _set_element_text(new_block[0], cap["display_name"])  # 改 heading2 标题
        fill_function_block_tables(new_block, cap)
        for el in new_block:
            anchor.addnext(el)
            anchor = el  # 锚点滚动，实现顺序追加

    # F. 非功能区无内容子章节插占位
    fill_nonfunctional_placeholders(doc)

    print(
        f"      映射：{len(capabilities)} capability → {len(capabilities)} 功能块 | "
        f"删除 {deleted} 个模板示例元素 | 功能块元素 {len(block_template)}/块"
    )


def enable_toc_auto_update(doc: Document) -> None:
    """设置 settings 让 Word 打开时自动更新 TOC 字段。"""
    settings = doc.settings.element
    # 避免重复添加
    for uf in settings.findall(qn("w:updateFields")):
        settings.remove(uf)
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


# ============================================================
# 主流程
# ============================================================

def generate(
    slug: str, project_arg: str | None, stage: str,
    cap_names_arg: str | None,
) -> Path:
    project_root = Path.cwd()
    check_pandoc()
    check_python_deps()

    change_dir = resolve_change_dir(slug, project_root)
    meta = read_meta(change_dir)
    template_path = find_template(change_dir, project_root)
    cap_names = load_capability_names(project_root, cap_names_arg)

    print(f"[1/5] 读取并解析 spec 产物：{change_dir}")
    proposal_md = (change_dir / "proposal.md").read_text(encoding="utf-8") \
        if (change_dir / "proposal.md").is_file() else ""
    proposal = parse_proposal(proposal_md)
    capabilities = parse_specs(change_dir)
    if not proposal_md and not capabilities:
        raise SddDocError(
            f"在 {change_dir} 下未找到 proposal.md 或 specs/*.md，无法生成正文。"
        )

    print(f"[2/5] 抽取元数据")
    data = extract_metadata(change_dir, project_root, meta, project_arg)
    revisions = extract_revisions(change_dir, project_root, meta)
    print(
        f"      项目：{data['project_title']} | 版本：{data['version']} | "
        f"修订行：{len(revisions)} | capability：{len(capabilities)}"
    )

    print(f"[3/5] python-docx 填充模板（套用：{template_path.name}）")
    out_path = change_dir / f"{slug}-{OUTPUT_FILENAME_SUFFIX}"
    shutil.copy(template_path, out_path)
    os.chmod(out_path, 0o644)

    doc = Document(str(out_path))
    fill_cover(doc, data)
    fill_info_table(doc, data)
    fill_revisions_table(doc, revisions)

    print(f"[4/5] 章节映射填充")
    map_to_template(doc, proposal, capabilities, cap_names)

    enable_toc_auto_update(doc)
    doc.save(str(out_path))

    print(f"[5/5] 完成：{out_path}")
    return out_path


def main() -> int:
    parser = argparse.ArgumentParser(
        prog="sdd_doc.py",
        description="从 openspec spec 产物生成企业格式 Word 文档",
    )
    parser.add_argument(
        "slug", nargs="?", default=None,
        help="change slug（openspec/changes/<slug>）；缺省时取唯一未归档 change",
    )
    parser.add_argument("--project", default=None, help="封面项目标题（覆盖 meta）")
    parser.add_argument(
        "--capability-names", default=None,
        help='capability 显示名映射，格式 "auth=认证,user=用户管理"（逗号分隔）',
    )
    parser.add_argument("--stage", default="spec", help="生成阶段（预留，当前仅 spec）")
    args = parser.parse_args()

    # slug 缺省：尝试找唯一未归档 change
    slug = args.slug
    if slug is None:
        changes_dir = Path.cwd() / "openspec" / "changes"
        if changes_dir.is_dir():
            candidates = [d.name for d in changes_dir.iterdir() if d.is_dir()]
            if len(candidates) == 1:
                slug = candidates[0]
            elif len(candidates) > 1:
                raise SddDocError(
                    f"存在多个 change，请指定 slug：{candidates}"
                )
        if slug is None:
            raise SddDocError("未指定 slug，且无法自动推断（openspec/changes/ 为空或不存在）。")

    try:
        out = generate(slug, args.project, args.stage, args.capability_names)
    except SddDocError as e:
        print(f"❌ {e}", file=sys.stderr)
        return 1
    print(f"\n✓ 已生成：{out}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
